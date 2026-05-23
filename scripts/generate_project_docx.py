"""
One-off script: generates college-level project documentation as .docx
Run from repo root: python scripts/generate_project_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(cell)
    doc.add_paragraph()


def main():
    root = Path(__file__).resolve().parent.parent
    out_path = root / "AI-Enhanced-Hospital-Queue-Project-Documentation.docx"

    doc = Document()

    title = doc.add_heading("AI-Enhanced Hospital Queue Management System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph("Project Documentation (College-Level)")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in st.runs:
        run.font.size = Pt(12)
    doc.add_paragraph()

    # --- 1 ---
    add_heading(doc, "1. Project Overview", 1)
    doc.add_paragraph(
        "This project is a full-stack hospital outpatient queue management system that combines "
        "a web application, a REST API, and a separate machine learning microservice. Patients "
        "can book time-slot–based appointments; on arrival, staff or patients perform check-in, "
        "which places the visit into a doctor- and slot-specific queue with tokens, priority ordering, "
        "and support for skip and no-show handling. The user interface receives live queue updates "
        "over WebSockets so doctors, receptionists, and patients see the current line without manual "
        "refresh. An AI-assisted layer estimates wait times and crowd levels to support scheduling "
        "and dashboard analytics."
    )
    doc.add_paragraph(
        "The system is structured as three runnable components: a Node.js/Express backend "
        "(port 5000 by default), a React (Vite) frontend (port 5173), and a Python FastAPI ML service "
        "(port 8001), with MongoDB as the primary data store."
    )

    # --- 2 ---
    add_heading(doc, "2. Problem Statement", 1)
    doc.add_paragraph(
        "Traditional hospital queues are often opaque and stressful: patients do not know their real "
        "position or wait time, staff manually call the next person with little systematic support "
        "for priority cases or absent patients, and peak-hour crowding is hard to anticipate. "
        "Paper tokens or simple first-come lines do not scale when departments mix urgent and routine "
        "visits or when no-shows waste slots."
    )
    doc.add_paragraph(
        "This project addresses these gaps by providing a digital, slot-based queue with role-based "
        "workflows, auditable events, real-time visibility, and data-driven estimates of wait and "
        "crowd intensity—suitable as a prototype for smarter outpatient flow in a teaching or demo environment."
    )

    # --- 3 ---
    add_heading(doc, "3. Objectives", 1)
    for item in [
        "Digitize appointment booking with clear slot windows per doctor and optional priority metadata.",
        "Model the live queue per doctor and per slot, with tokens, ordering, and status transitions.",
        "Expose a secure API with authentication, validation, and role checks for sensitive operations.",
        "Push real-time updates to connected clients when the queue changes.",
        "Integrate an ML service for wait-time and crowd predictions, with explanations suitable for dashboards.",
        "Provide a modern web UI for login, booking, queue view, appointments, and analytics-oriented screens.",
    ]:
        doc.add_paragraph(item, style="List Number")

    # --- 4 ---
    add_heading(doc, "4. Key Features (with Explanation)", 1)
    add_table(
        doc,
        ["Feature", "Explanation"],
        [
            (
                "JWT authentication",
                "Users sign up and log in; the API returns a JSON Web Token. Protected routes verify the token.",
            ),
            (
                "Role-based users",
                "Roles: Admin, Receptionist, Doctor, Patient. Non-patient signups may require admin activation (isActive).",
            ),
            (
                "Appointment booking",
                "Book slot with doctorId, slotStartAt, slotEndAt; optional reason and priority (EMERGENCY, SENIOR, PREGNANT, VIP, NORMAL).",
            ),
            (
                "Check-in → queue entry",
                "Check-in creates or reuses a Queue for that doctor and slot; embeds entry with token and position.",
            ),
            (
                "Token + priority ordering",
                "Next patient: WAITING entries sorted by higher priorityScore first, then lower position.",
            ),
            (
                "Call next / skip / no-show",
                "Staff advance queue, skip entries, or mark no-show; MongoDB updates and Socket.io broadcasts.",
            ),
            (
                "Audit trail",
                "QueueLog records events (e.g. CHECK_IN, CALL_NEXT) with actor and metadata.",
            ),
            (
                "Live queue & activity feed",
                "Clients join room queue:doctorId:slotISO; server emits queue:updated. Global activity:event for highlights.",
            ),
            (
                "ML wait & crowd prediction",
                "Backend proxies to FastAPI: wait bands + explanations; crowd level and arrivals per hour.",
            ),
            (
                "Doctor directory",
                "GET /api/doctors lists doctors for booking.",
            ),
            (
                "Admin approval",
                "Endpoints to approve doctors and users.",
            ),
            (
                "Dashboard & booking UI",
                "Dashboard uses predictions for charts; booking flow can show ML wait hints.",
            ),
        ],
    )
    doc.add_paragraph(
        "Note: Some UI areas (e.g. Statement billing) use static prototype data with a note to connect the backend later."
    )

    # --- 5 ---
    add_heading(doc, "5. Complete Tech Stack (Role of Each Technology)", 1)
    add_heading(doc, "Frontend", 2)
    for line in [
        "React 19 — UI components and state.",
        "Vite — Dev server and production build.",
        "React Router — Client-side routing.",
        "Tailwind CSS — Styling.",
        "socket.io-client — WebSocket client for real-time events.",
        "Recharts — Dashboard charts.",
        "Lucide React — Icons.",
        "clsx — Conditional class names.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "Backend (API + Real-Time)", 2)
    for line in [
        "Node.js — Runtime.",
        "Express 5 — HTTP routing and middleware.",
        "Mongoose — MongoDB ODM.",
        "MongoDB — Primary database.",
        "Socket.io — WebSocket server on the same HTTP server.",
        "jsonwebtoken — JWT creation and verification.",
        "bcrypt — Password hashing.",
        "Zod — Request validation.",
        "helmet — Security-related HTTP headers.",
        "cors — Cross-origin access control.",
        "express-rate-limit — Rate limiting.",
        "morgan — HTTP logging.",
        "dotenv — Environment configuration.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "ML Service", 2)
    for line in [
        "Python 3 — ML service runtime.",
        "FastAPI — REST API for predictions.",
        "Uvicorn — ASGI server.",
        "NumPy — Arrays, synthetic data, ridge regression.",
        "Pydantic — Request/response schemas.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "DevOps / Tooling", 2)
    for line in [
        "concurrently — Run backend, frontend, and ML together (npm run dev).",
        "nodemon — Backend auto-restart in development.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # --- 6 ---
    add_heading(doc, "6. System Architecture and Workflow", 1)
    doc.add_paragraph(
        "High-level flow: (1) The browser loads the React app and uses REST + Socket.io to the Express server. "
        "(2) Express reads/writes MongoDB via Mongoose. (3) For predictions, Express calls FastAPI over HTTP. "
        "(4) On queue changes, Express emits Socket.io events to subscribed clients."
    )
    doc.add_paragraph(
        "Typical patient journey: Register/login → receive JWT → book appointment (BOOKED) → check-in → "
        "appointment becomes IN_QUEUE with a queue entry → staff call next (CALLED) → optional skip or no-show; "
        "clients see updates live via sockets."
    )
    doc.add_paragraph(
        "Conceptual diagram (text): Browser (React) ↔ Express (REST + Socket.io) ↔ MongoDB; Express ↔ FastAPI (ML)."
    )

    # --- 7 ---
    add_heading(doc, "7. Database Design", 1)
    doc.add_paragraph(
        "MongoDB stores documents. Main collections:"
    )
    add_heading(doc, "Users (User)", 2)
    doc.add_paragraph(
        "name, email (unique), optional phone (sparse unique), role, passwordHash, isActive, lastLoginAt; "
        "bcrypt verify/hash methods."
    )
    add_heading(doc, "Doctors (Doctor)", 2)
    doc.add_paragraph(
        "Optional userId; name, department, specialization, registrationNo; slotTemplates (weekday, times, "
        "slotMinutes, capacity); avgServiceMinutes; isActive; text index for search."
    )
    add_heading(doc, "Appointments (Appointment)", 2)
    doc.add_paragraph(
        "patientId, doctorId, reason, slotStartAt/slotEndAt, priorityLevel, priorityScore, status lifecycle, "
        "timestamps, queueId, queueEntryId; indexes for doctor+slot and patient history."
    )
    add_heading(doc, "Queues (Queue)", 2)
    doc.add_paragraph(
        "Unique per (doctorId, slotStartAt). Fields: slotEndAt, status, currentTokenNo, nextTokenNo, "
        "embedded entries array (appointmentId, patientId, priority, status, tokenNo, position, etc.). "
        "Embedded entries enable fast reads; QueueLog provides audit."
    )
    add_heading(doc, "QueueLogs (QueueLog)", 2)
    doc.add_paragraph(
        "queueId, doctorId, optional appointment/entry/patient refs, eventType, from/to status, message, meta, actorUserId."
    )

    # --- 8 ---
    add_heading(doc, "8. API Explanation", 1)
    add_table(
        doc,
        ["Area", "Method & Path", "Purpose"],
        [
            ("Health", "GET /health", "Liveness check."),
            ("Auth", "POST /api/auth/signup", "Create user; returns JWT."),
            ("Auth", "POST /api/auth/login", "Login; returns JWT."),
            ("Auth", "GET /api/auth/me", "Current user (Bearer token)."),
            ("Doctors", "GET /api/doctors", "List doctors."),
            ("Appointments", "POST /api/appointments/book", "Book slot (auth)."),
            ("Appointments", "POST /api/appointments/check-in", "Check-in; add to queue (role rules)."),
            ("Appointments", "GET /api/appointments/mine", "List appointments for user."),
            ("Queues", "GET /api/queues?doctorId=&slotStartAt=", "Get formatted queue."),
            ("Queues", "POST /api/queues/next", "Call next (doctor/receptionist/admin)."),
            ("Queues", "POST /api/queues/no-show", "Mark no-show."),
            ("Queues", "POST /api/queues/skip", "Skip entry."),
            ("Admin", "POST /api/admin/approve-doctor", "Approve doctor active flag."),
            ("Admin", "POST /api/admin/approve-user", "Approve user active flag."),
            ("Predict", "POST /api/predict/wait-time", "Proxy to ML wait prediction."),
            ("Predict", "POST /api/predict/crowd", "Proxy to ML crowd prediction."),
        ],
    )
    doc.add_paragraph("Responses typically use { ok: true, data: ... } with centralized error handling.")

    # --- 9 ---
    add_heading(doc, "9. AI/ML Module Explanation", 1)
    doc.add_paragraph(
        "Location: ml-service (FastAPI). On startup, synthetic data is generated with department and time patterns; "
        "a feature matrix is built (queue length, arrivals, service rate, emergency share, sin/cos hour/DOW, "
        "one-hot department). Two ridge-style linear models are trained: wait minutes and crowd (arrivals per hour)."
    )
    doc.add_paragraph(
        "Endpoints: POST /predict/wait-time returns wait_minutes, band (low/moderate/high), explanation, top_factors. "
        "POST /predict/crowd returns crowd_level, predicted_arrivals_per_hour, explanation."
    )
    doc.add_paragraph(
        "The Node backend calls the ML service over HTTP (no embedded model). Department names are normalized in the backend. "
        "For academic honesty: models use synthetic training data; production would use real de-identified data and validation."
    )

    # --- 10 ---
    add_heading(doc, "10. Real-Time Communication Flow", 1)
    for item in [
        "HTTP server + Socket.io with CORS matching client origin.",
        "On connect: socket joins activity:global.",
        "Client emits join { doctorId, slotStartAt }; server joins room queue:{doctorId}:{ISO slot}.",
        "On queue change: emit queue:updated with full queue payload.",
        "Activity: emit activity:event with type (e.g. PATIENT_CHECKED_IN, TOKEN_CALLED).",
        "Frontend: useQueueLive fetches REST then listens for queue:updated; useActivityFeed listens for activity:event.",
        "ISO strings for slotStartAt keep room names consistent.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # --- 11 ---
    add_heading(doc, "11. Security Features", 1)
    for item in [
        "Passwords: bcrypt hashing.",
        "JWT authentication on protected routes.",
        "Authorization: requireRole for check-in, queue ops, admin; patients only check in own appointments.",
        "HTTP: Helmet, JSON body limit (1 MB), rate limiting.",
        "CORS: configured origin with credentials.",
        "Input validation: Zod on routes.",
        "Account gating: non-patient signups may start inactive until admin approval.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Note for deployment: /api/predict/* routes are validated but not JWT-protected in code—consider auth or "
        "stricter gateway rules in production."
    )

    # --- 12 ---
    add_heading(doc, "12. Challenges Faced During Development", 1)
    for item in [
        "Multi-service coordination (API, UI, ML) and environment variables.",
        "Real-time consistency: matching room keys (doctor id + slot ISO) across clients.",
        "Embedded vs normalized data: fast queue reads vs audit in QueueLog.",
        "Priority semantics: deterministic next-patient rule (priorityScore then position).",
        "ML service availability: graceful degradation when FastAPI is down.",
        "Role and workflow edge cases: idempotent check-in, valid status transitions, booking on behalf of patients.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # --- 13 ---
    add_heading(doc, "13. Future Scope", 1)
    for item in [
        "Production ML on privacy-compliant real logs; confidence intervals; per-department calibration.",
        "Notifications: SMS/email/push when token is next or called.",
        "Billing: connect Statement page to payments.",
        "Scale: Redis for Socket.io; job queues for analytics.",
        "Mobile app or reception kiosk.",
        "FHIR/EHR interoperability.",
        "Stronger auth: refresh tokens, OAuth2, MFA for staff.",
        "Automated tests (API, sockets, ML).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # --- 14 ---
    add_heading(doc, "14. Conclusion", 1)
    doc.add_paragraph(
        "The AI-Enhanced Hospital Queue Management System delivers an integrated prototype combining structured "
        "appointment and queue management, role-based access, real-time updates, and machine-learning-assisted wait "
        "and crowd insights. The stack—React, Node/Express, MongoDB, Socket.io, and FastAPI/NumPy—shows how modern "
        "web, data, and lightweight ML services compose for healthcare operations education. While the ML layer uses "
        "synthetic data and some UI modules are demonstration-only, the project demonstrates end-to-end flow from "
        "booking to live queue control and analytics, with a clear path to production hardening."
    )

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
