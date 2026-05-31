"""
Generate printable Viva Prep outputs from docs/VIVA_PREP.md

Run from repo root:
  python scripts/generate_viva_prep.py

Outputs:
  docs/VIVA_PREP.html  — open in browser, Print → Save as PDF
  docs/VIVA_PREP.docx  — open in Word, Save as PDF (requires python-docx)
"""
from __future__ import annotations

import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "VIVA_PREP.md"
HTML_PATH = ROOT / "docs" / "VIVA_PREP.html"
DOCX_PATH = ROOT / "docs" / "VIVA_PREP.docx"

PRINT_CSS = """
@page { size: A4; margin: 18mm 16mm 20mm 16mm; }
* { box-sizing: border-box; }
body {
  font-family: "Segoe UI", Calibri, Arial, sans-serif;
  font-size: 11pt;
  line-height: 1.45;
  color: #1a1a1a;
  max-width: 210mm;
  margin: 0 auto;
  padding: 12mm 0;
}
.cover {
  text-align: center;
  padding: 24mm 8mm 16mm;
  page-break-after: always;
  border-bottom: 3px solid #0d47a1;
}
.cover h1 {
  font-size: 22pt;
  color: #0d47a1;
  margin: 0 0 8px;
  line-height: 1.2;
}
.cover .subtitle { font-size: 13pt; color: #37474f; margin: 0 0 20px; }
.cover .meta { font-size: 10pt; color: #546e7a; }
.cover .stack { margin-top: 16px; font-size: 9.5pt; line-height: 1.6; }
h2 {
  font-size: 14pt;
  color: #0d47a1;
  margin: 22px 0 10px;
  padding-bottom: 4px;
  border-bottom: 1px solid #b0bec5;
  page-break-after: avoid;
}
h3 {
  font-size: 12pt;
  color: #1565c0;
  margin: 16px 0 8px;
  page-break-after: avoid;
}
h4 {
  font-size: 11pt;
  color: #263238;
  margin: 14px 0 6px;
  page-break-after: avoid;
}
p { margin: 0 0 8px; }
ul, ol { margin: 0 0 10px; padding-left: 22px; }
li { margin-bottom: 4px; }
hr { border: none; border-top: 1px solid #cfd8dc; margin: 16px 0; }
blockquote {
  margin: 10px 0 14px;
  padding: 10px 14px;
  background: #e3f2fd;
  border-left: 4px solid #1976d2;
  font-size: 10.5pt;
  page-break-inside: avoid;
}
table {
  width: 100%;
  border-collapse: collapse;
  margin: 10px 0 14px;
  font-size: 10pt;
  page-break-inside: avoid;
}
th, td {
  border: 1px solid #90a4ae;
  padding: 6px 8px;
  text-align: left;
  vertical-align: top;
}
th { background: #eceff1; font-weight: 600; }
.qa-block {
  margin-bottom: 12px;
  page-break-inside: avoid;
}
.toc { font-size: 10.5pt; }
.toc ol { padding-left: 20px; }
.toc a { color: #1565c0; text-decoration: none; }
.footer-note {
  margin-top: 24px;
  padding-top: 10px;
  border-top: 1px solid #cfd8dc;
  font-size: 9pt;
  color: #607d8b;
  font-style: italic;
}
@media print {
  body { padding: 0; }
  a { color: inherit; text-decoration: none; }
  .no-print { display: none !important; }
  h2, h3, h4 { page-break-after: avoid; }
  .qa-block, blockquote, table { page-break-inside: avoid; }
}
.print-hint {
  background: #fff8e1;
  border: 1px solid #ffc107;
  padding: 10px 14px;
  margin-bottom: 20px;
  border-radius: 4px;
  font-size: 10pt;
}
"""


def escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def inline_format(text: str) -> str:
    text = escape_html(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    return text


def parse_table(lines: list[str]) -> str:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return ""
    html = ["<table>"]
    for i, row in enumerate(rows):
        tag = "th" if i == 0 else "td"
        html.append("<tr>" + "".join(f"<{tag}>{inline_format(c)}</{tag}>" for c in row) + "</tr>")
    html.append("</table>")
    return "\n".join(html)


def md_to_html_body(md: str) -> str:
    lines = md.splitlines()
    out: list[str] = []
    i = 0
    in_blockquote = False
    bq_lines: list[str] = []
    list_items: list[str] = []
    list_ordered = False
    table_buf: list[str] = []
    skip_until_part1 = True

    def flush_list():
        nonlocal list_items, list_ordered
        if not list_items:
            return
        tag = "ol" if list_ordered else "ul"
        out.append(f"<{tag}>")
        for item in list_items:
            out.append(f"<li>{inline_format(item)}</li>")
        out.append(f"</{tag}>")
        list_items = []

    def flush_bq():
        nonlocal in_blockquote, bq_lines
        if bq_lines:
            out.append("<blockquote>" + "<br/>".join(inline_format(l) for l in bq_lines) + "</blockquote>")
        bq_lines = []
        in_blockquote = False

    def flush_table():
        nonlocal table_buf
        if table_buf:
            out.append(parse_table(table_buf))
            table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if skip_until_part1:
            if stripped.startswith("## Part 1"):
                skip_until_part1 = False
            else:
                i += 1
                continue

        if stripped.startswith("|"):
            flush_list()
            flush_bq()
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if stripped.startswith("> "):
            flush_list()
            if not in_blockquote:
                in_blockquote = True
            bq_lines.append(stripped[2:].strip())
            i += 1
            continue
        flush_bq()

        if re.match(r"^[-*] \[ \] ", stripped):
            flush_list()
            out.append(f'<p class="checklist">☐ {inline_format(stripped[6:])}</p>')
            i += 1
            continue

        if re.match(r"^[-*] ", stripped):
            if list_ordered:
                flush_list()
            list_items.append(stripped[2:].strip())
            list_ordered = False
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            if not list_ordered and list_items:
                flush_list()
            list_items.append(re.sub(r"^\d+\.\s*", "", stripped))
            list_ordered = True
            i += 1
            continue

        flush_list()

        if stripped == "---":
            out.append("<hr/>")
            i += 1
            continue

        if stripped.startswith("#### "):
            title = stripped[5:].strip()
            if title.startswith("Q"):
                out.append(f'<div class="qa-block"><h4>{inline_format(title)}</h4>')
            else:
                out.append(f"<h4>{inline_format(title)}</h4>")
            i += 1
            continue

        if stripped.startswith("### "):
            out.append(f"<h3>{inline_format(stripped[4:])}</h3>")
            i += 1
            continue

        if stripped.startswith("## "):
            out.append(f"<h2 id='{slug(stripped[3:])}'>{inline_format(stripped[3:])}</h2>")
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue

        if not stripped:
            if out and out[-1].startswith("<p>"):
                out.append("</div>")
            i += 1
            continue

        if out and out[-1].startswith('<div class="qa-block"><h4>'):
            if not (out[-1].endswith("</h4>") or "</h4>" in out[-1]):
                pass
            # Answer paragraph after question
            out.append(f"<p>{inline_format(stripped)}</p>")
        else:
            out.append(f"<p>{inline_format(stripped)}</p>")
        i += 1

    flush_list()
    flush_bq()
    flush_table()
    html = "\n".join(out)
    html = re.sub(r'(<div class="qa-block"><h4>.*?</h4>)\s*(<p>)', r"\1\2", html, flags=re.DOTALL)
    # Close open qa-blocks before next h4 or h3
    parts = re.split(r'(?=<div class="qa-block">)', html)
    fixed = []
    for idx, part in enumerate(parts):
        if 'class="qa-block"' in part and part.count("<div") > part.count("</div>"):
            part = part + "</div>"
        fixed.append(part)
    return "".join(fixed)


def slug(text: str) -> str:
    s = text.lower()
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s[:60]


def build_html(md: str) -> str:
    body = md_to_html_body(md)
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <title>Viva Prep — AI-Enhanced Hospital Queue Management System</title>
  <style>{PRINT_CSS}</style>
</head>
<body>
  <p class="print-hint no-print"><strong>Print to PDF:</strong> Press <kbd>Ctrl+P</kbd> (or <kbd>Cmd+P</kbd> on Mac) → choose <em>Save as PDF</em> → enable <em>Background graphics</em> for best results.</p>

  <section class="cover">
    <h1>AI-Enhanced Hospital Queue Management System</h1>
    <p class="subtitle">Viva &amp; Interview Preparation Guide</p>
    <p class="meta">38 Questions · Top 10 Dangerous Questions · 2-Minute Pitch · Architecture · Common Mistakes</p>
    <p class="stack">
      <strong>Stack:</strong> React (Vite) · Node.js + Express · MongoDB Atlas · Socket.io · FastAPI ML · Tailwind · JWT<br/>
      <strong>Deploy:</strong> Vercel · Render · MongoDB Atlas
    </p>
  </section>

  <nav class="toc">
    <h2>Table of Contents</h2>
    <ol>
      <li>Part 1 — Viva Questions &amp; Answers (38)</li>
      <li>Part 2 — Top 10 Most Dangerous Viva Questions</li>
      <li>Part 3 — 2-Minute Project Explanation</li>
      <li>Part 4 — Architecture Diagram Explanation</li>
      <li>Part 5 — Common Mistakes During Viva</li>
      <li>Quick Revision Checklist &amp; Demo Accounts</li>
    </ol>
  </nav>

  <hr/>

  {body}

  <p class="footer-note">Project viva preparation document. Regenerate with: python scripts/generate_viva_prep.py</p>
</body>
</html>
"""


def build_docx(md: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise SystemExit("python-docx not installed. Run: pip install python-docx") from e

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_heading("AI-Enhanced Hospital Queue Management System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Viva & Interview Preparation Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)
    doc.add_paragraph()

    in_part = False
    i = 0
    lines = md.splitlines()
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not in_part:
            if stripped.startswith("## Part 1"):
                in_part = True
            else:
                i += 1
                continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[5:])
            run.bold = True
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                if all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                rows.append(cells)
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        t.rows[ri].cells[ci].text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell)
                doc.add_paragraph()
            continue
        if stripped.startswith("> "):
            doc.add_paragraph(stripped[2:], style="Intense Quote")
            i += 1
            continue
        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("- [ ]"):
            doc.add_paragraph("☐ " + stripped[6:], style="List Bullet")
            i += 1
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        text = re.sub(r"`(.+?)`", r"\1", text)
        doc.add_paragraph(text)
        i += 1

    doc.save(DOCX_PATH)


def main():
    if not MD_PATH.exists():
        raise SystemExit(f"Missing source file: {MD_PATH}")

    md = MD_PATH.read_text(encoding="utf-8")
    html = build_html(md)
    HTML_PATH.write_text(html, encoding="utf-8")
    print(f"Wrote {HTML_PATH.relative_to(ROOT)}")
    print("  -> Open in Chrome/Edge -> Print -> Save as PDF")

    try:
        build_docx(md)
        print(f"Wrote {DOCX_PATH.relative_to(ROOT)}")
        print("  -> Open in Word -> File -> Save As -> PDF")
    except SystemExit as e:
        print(e)


if __name__ == "__main__":
    main()
